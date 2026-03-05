"""
Export resume sections back to .docx format.

Converts section HTML back to python-docx paragraphs preserving
bold, italic, underline, alignment, color, and bullet formatting.
Supports adjustable margins and right-aligned date tab stops.
"""

import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MARGIN_PRESETS = {
    'narrow': (0.4, 0.4, 0.4, 0.4),
    'normal': (0.5, 0.5, 0.5, 0.5),
    'wide': (0.75, 0.75, 0.75, 0.75),
}


class _HTMLToRuns(HTMLParser):
    """Parse HTML into a list of (text, bold, italic, underline, color, float_right, href) tuples."""

    def __init__(self):
        super().__init__()
        self.runs = []
        self._bold = False
        self._italic = False
        self._underline = False
        self._color = None
        self._float_right = False
        self._href = None
        self._tag_stack = []

    def handle_starttag(self, tag, attrs):
        self._tag_stack.append(tag)
        if tag in ('strong', 'b'):
            self._bold = True
        elif tag in ('em', 'i'):
            self._italic = True
        elif tag == 'u':
            self._underline = True
        elif tag == 'a':
            attrs_dict = dict(attrs)
            self._href = attrs_dict.get('href')
        elif tag == 'span':
            attrs_dict = dict(attrs)
            style = attrs_dict.get('style', '')
            color_match = re.search(r'color:\s*#([0-9a-fA-F]{6})', style)
            if color_match:
                self._color = color_match.group(1)
            else:
                # Handle rgb(r, g, b) format from rich text editors
                rgb_match = re.search(r'color:\s*rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', style)
                if rgb_match:
                    r, g, b = int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3))
                    self._color = f'{r:02x}{g:02x}{b:02x}'
            if 'float:right' in style or 'float: right' in style:
                self._float_right = True

    def handle_endtag(self, tag):
        if tag in ('strong', 'b'):
            self._bold = False
        elif tag in ('em', 'i'):
            self._italic = False
        elif tag == 'u':
            self._underline = False
        elif tag == 'a':
            self._href = None
        elif tag == 'span':
            self._color = None
            self._float_right = False
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

    def handle_data(self, data):
        if data:
            self.runs.append((data, self._bold, self._italic, self._underline, self._color, self._float_right, self._href))

    def handle_entityref(self, name):
        entities = {'amp': '&', 'lt': '<', 'gt': '>', 'quot': '"',
                    'emsp': '\t', 'nbsp': '\u00a0'}
        char = entities.get(name, f'&{name};')
        self.runs.append((char, self._bold, self._italic, self._underline, self._color, self._float_right, self._href))

    def handle_charref(self, name):
        try:
            char = chr(int(name, 16) if name.startswith('x') else int(name))
        except ValueError:
            char = f'&#{name};'
        self.runs.append((char, self._bold, self._italic, self._underline, self._color, self._float_right, self._href))


def _parse_inline_html(html):
    """Parse inline HTML to a list of run tuples."""
    parser = _HTMLToRuns()
    parser.feed(html)
    return parser.runs


# Date pattern for auto-detecting right-aligned dates
_DATE_RE = re.compile(
    r'(\d{1,2}/)?\d{4}\s*[–\-—]\s*(\d{1,2}/)?\d{4}|(\d{1,2}/)?\d{4}\s*[–\-—]\s*[Pp]resent'
)


def _auto_float_dates(runs_data):
    """Detect date patterns in runs and mark them as float_right for tab alignment.

    Handles editor-saved HTML where float:right spans were stripped.
    """
    # Skip if already has float_right runs
    if any(f for _, _, _, _, _, f, _ in runs_data):
        return runs_data

    # Concatenate all text to find date pattern
    full_text = ''.join(t for t, *_ in runs_data)
    date_match = _DATE_RE.search(full_text)
    if not date_match:
        return runs_data

    date_start = date_match.start()

    # Find which run contains the date start and split there
    new_runs = []
    char_pos = 0
    found = False
    for text, bold, italic, underline, color, float_right, href in runs_data:
        if found:
            new_runs.append((text, bold, italic, underline, color, True, href))
        elif char_pos + len(text) > date_start:
            split_at = date_start - char_pos
            before = text[:split_at]
            after = text[split_at:]
            if before.strip():
                new_runs.append((before, bold, italic, underline, color, False, href))
            if after:
                new_runs.append((after, bold, italic, underline, color, True, href))
            found = True
        else:
            new_runs.append((text, bold, italic, underline, color, False, href))
        char_pos += len(text)

    return new_runs if found else runs_data


def _detect_alignment(tag_html):
    """Extract text-align from an element's style attribute."""
    m = re.search(r'text-align:\s*(right|center|left)', tag_html)
    if m:
        val = m.group(1)
        if val == 'right':
            return WD_ALIGN_PARAGRAPH.RIGHT
        elif val == 'center':
            return WD_ALIGN_PARAGRAPH.CENTER
    return None


def _detect_font_size(attrs_html):
    """Extract font size (pt) from data-font-size attribute."""
    m = re.search(r'data-font-size="(\d+(?:\.\d+)?)"', attrs_html)
    if m:
        return float(m.group(1))
    return None


def _detect_borders(attrs_html):
    """Extract border info from data-border-* attributes."""
    borders = {}
    for side in ('bottom', 'top'):
        m = re.search(rf'data-border-{side}="([^"]+)"', attrs_html)
        if m:
            parts = m.group(1).split(',')
            borders[side] = {
                'val': parts[0] if len(parts) > 0 else 'single',
                'sz': parts[1] if len(parts) > 1 else '4',
                'color': parts[2] if len(parts) > 2 else 'auto',
            }
    return borders


# Regex to split HTML into block elements
_BLOCK_RE = re.compile(
    r'<(p|li|h[23]|ul|/ul)(\s[^>]*)?>',
    re.IGNORECASE,
)


def _split_blocks(html):
    """Split HTML into a list of (tag, attrs_html, inner_html) blocks."""
    blocks = []
    current_tag = None
    current_attrs = ''
    content_start = 0
    # Track if we're inside a <li> so nested <p> doesn't override it
    in_li = False

    for m in re.finditer(r'<(/?)(\w+)([^>]*)>', html):
        is_close = m.group(1) == '/'
        tag = m.group(2).lower()
        attrs = m.group(3)

        if tag in ('ul',):
            continue  # skip ul wrappers

        if tag in ('p', 'li', 'h2', 'h3'):
            if is_close:
                if tag == 'p' and in_li:
                    continue  # skip </p> inside <li>
                if tag == 'li':
                    in_li = False
                if current_tag:
                    inner = html[content_start:m.start()]
                    blocks.append((current_tag, current_attrs, inner))
                    current_tag = None
            else:
                if tag == 'p' and in_li:
                    # Skip <p> inside <li> — use the <li> as the block tag
                    content_start = m.end()
                    continue
                if tag == 'li':
                    in_li = True
                current_tag = tag
                current_attrs = attrs
                content_start = m.end()

    return blocks


class ResumeExporter:
    """Export resume sections to a .docx file."""

    def export(self, sections, output_path, margin='normal'):
        """Generate .docx from section dicts with content_html."""
        doc = Document()

        # Set margins based on preset
        margins = MARGIN_PRESETS.get(margin, MARGIN_PRESETS['normal'])
        top, bottom, left, right = margins
        for section in doc.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)

        # Calculate content width for right-aligned tab stops
        self._content_width = Inches(8.5 - left - right)

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.space_before = Pt(0)

        for sec in sections:
            html = sec.get('content_html', sec.get('html', ''))
            if not html.strip():
                continue
            sec_type = sec.get('section_type', sec.get('type', ''))
            self._write_section(doc, html, sec_type)

        doc.save(output_path)

    def _write_section(self, doc, html, section_type):
        """Write a single section's HTML content to the document."""
        blocks = _split_blocks(html)

        # Pre-compute indices of <p> blocks for summary border logic
        p_indices = [i for i, (tag, _, _) in enumerate(blocks) if tag == 'p']

        for block_idx, (tag, attrs, inner) in enumerate(blocks):
            alignment = _detect_alignment(attrs)
            runs_data = _parse_inline_html(inner)

            # Auto-detect dates for right-alignment (not for bullets)
            if tag != 'li':
                runs_data = _auto_float_dates(runs_data)

            if tag in ('h2', 'h3'):
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(2)
                if alignment:
                    para.alignment = alignment
                font_size = _detect_font_size(attrs) or 14
                # Use _add_runs so float_right dates get tab-aligned
                has_float_right = any(f for _, _, _, _, _, f, _ in runs_data)
                if has_float_right:
                    self._add_runs(doc, para, runs_data, font_size=font_size, force_bold=True)
                else:
                    for text, bold, italic, underline, color, float_right, href in runs_data:
                        if href:
                            self._add_hyperlink(doc, para, text, href, bold=True, italic=italic, underline=underline, color=color, font_size=Pt(font_size))
                        else:
                            run = para.add_run(text)
                            run.bold = True
                            run.italic = italic
                            run.underline = underline
                            run.font.size = Pt(font_size)
                            if color:
                                run.font.color.rgb = RGBColor.from_string(color)
                borders = _detect_borders(attrs)
                if not borders:
                    # Default bottom border for section headings
                    borders = {'bottom': {'val': 'single', 'sz': '4', 'color': 'auto'}}
                self._apply_borders(para, borders)

            elif tag == 'li':
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.space_before = Pt(0)
                if alignment:
                    para.alignment = alignment
                self._add_runs(doc, para, runs_data)

            else:  # <p>
                para = doc.add_paragraph()
                if alignment:
                    para.alignment = alignment

                font_size = _detect_font_size(attrs)
                borders = _detect_borders(attrs)

                if font_size and font_size > 10:
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(2)

                if section_type == 'header' and self._is_name_line(runs_data):
                    # Name line: larger font
                    for text, bold, italic, underline, color, float_right, href in runs_data:
                        if href:
                            self._add_hyperlink(doc, para, text, href, bold=bold, italic=italic, underline=underline, font_size=Pt(20))
                        else:
                            run = para.add_run(text)
                            run.bold = bold
                            run.italic = italic
                            run.underline = underline
                            run.font.size = Pt(20)
                            if color:
                                run.font.color.rgb = RGBColor.from_string(color)
                else:
                    self._add_runs(doc, para, runs_data, font_size=font_size)

                if not borders and section_type == 'summary' and len(p_indices) >= 2:
                    # Add bottom border after summary text and after Tools line
                    if block_idx == p_indices[0] or block_idx == p_indices[-1]:
                        borders = {'bottom': {'val': 'single', 'sz': '6', 'color': 'auto'}}
                        para.paragraph_format.space_after = Pt(6)

                if borders:
                    self._apply_borders(para, borders)

    def _add_runs(self, doc, para, runs_data, font_size=None, force_bold=False):
        """Add formatted runs to a paragraph, handling color, hyperlinks, and right-aligned dates."""
        has_float_right = any(f for t, b, i, u, c, f, h in runs_data)

        if has_float_right:
            # Add a right-aligned tab stop at the content width
            para.paragraph_format.tab_stops.add_tab_stop(
                self._content_width,
                alignment=WD_TAB_ALIGNMENT.RIGHT,
            )

        tab_inserted = False
        for text, bold, italic, underline, color, float_right, href in runs_data:
            if float_right and not tab_inserted:
                # Insert tab ONCE to jump to the right-aligned tab stop
                para.add_run('\t')
                tab_inserted = True

            effective_bold = bold or force_bold
            if href:
                self._add_hyperlink(doc, para, text, href, bold=effective_bold, italic=italic, underline=underline, color=color, font_size=Pt(font_size) if font_size else None)
            else:
                run = para.add_run(text)
                run.bold = effective_bold or None
                run.italic = italic or None
                run.underline = underline or None
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)
                if font_size:
                    run.font.size = Pt(font_size)

    def _add_hyperlink(self, doc, para, text, url, bold=False, italic=False, underline=False, color=None, font_size=None):
        """Add a clickable hyperlink to a paragraph."""
        part = para.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Blue color and underline for links (standard link styling)
        c_elem = OxmlElement('w:color')
        if color:
            c_elem.set(qn('w:val'), color)
        else:
            c_elem.set(qn('w:val'), '0563C1')
        rPr.append(c_elem)

        u_elem = OxmlElement('w:u')
        u_elem.set(qn('w:val'), 'single')
        rPr.append(u_elem)

        if bold:
            b_elem = OxmlElement('w:b')
            rPr.append(b_elem)
        if italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)
        if font_size:
            sz_elem = OxmlElement('w:sz')
            sz_elem.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz_elem)

        new_run.append(rPr)
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = text
        new_run.append(t_elem)
        hyperlink.append(new_run)
        para._element.append(hyperlink)

    def _apply_borders(self, para, borders):
        """Apply paragraph borders (top/bottom) from parsed border data."""
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side, props in borders.items():
            border_el = OxmlElement(f'w:{side}')
            border_el.set(qn('w:val'), props['val'])
            border_el.set(qn('w:sz'), props['sz'])
            border_el.set(qn('w:space'), '1')
            border_el.set(qn('w:color'), props['color'])
            pBdr.append(border_el)
        # Insert pBdr before spacing/tabs/jc per OpenXML element ordering
        pPr.insert(0, pBdr)

    def _is_name_line(self, runs_data):
        """Check if runs look like a name line (first paragraph, bold, short)."""
        if not runs_data:
            return False
        text = ''.join(t for t, b, i, u, c, f, h in runs_data).strip()
        all_bold = all(b for t, b, i, u, c, f, h in runs_data if t.strip())
        return all_bold and len(text) < 60 and '|' in text
